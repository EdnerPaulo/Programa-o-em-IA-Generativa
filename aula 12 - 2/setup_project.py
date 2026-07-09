import os

# Definição da estrutura de arquivos com os códigos fornecidos na arquitetura anterior
project_files = {
    "app/core/config.py": """from pydantic_settings import BaseSettings
from pydantic import ConfigDict

class Settings(BaseSettings):
    DATABASE_URL: str
    GEMINI_API_KEY: str
    JWT_SECRET_KEY: str
    JWT_ALGORITHM: str = "HS256"
    ACCESS_TOKEN_EXPIRE_MINUTES: int = 30
    REFRESH_TOKEN_EXPIRE_DAYS: int = 7
    MAX_FILE_SIZE_BYTES: int = 15 * 1024 * 1024  # 15MB
    UPLOAD_DIR: str = "./uploads"

    model_config = ConfigDict(env_file=".env", extra="ignore")

settings = Settings()""",

    "app/core/security.py": """from datetime import datetime, timedelta, timezone
from typing import Optional
from jose import jwt, JWTError
from passlib.context import CryptContext
from app.core.config import settings

pwd_context = CryptContext(schemes=["bcrypt"], deprecated="auto")

def get_password_hash(password: str) -> str:
    return pwd_context.hash(password)

def verify_password(plain_password: str, hashed_password: str) -> bool:
    return pwd_context.verify(plain_password, hashed_password)

def create_access_token(data: dict, expires_delta: Optional[timedelta] = None) -> str:
    to_encode = data.copy()
    expire = datetime.now(timezone.utc) + (expires_delta or timedelta(minutes=settings.ACCESS_TOKEN_EXPIRE_MINUTES))
    to_encode.update({"exp": expire})
    return jwt.encode(to_encode, settings.JWT_SECRET_KEY, algorithm=settings.JWT_ALGORITHM)

def create_refresh_token(data: dict, expires_delta: Optional[timedelta] = None) -> str:
    to_encode = data.copy()
    expire = datetime.now(timezone.utc) + (expires_delta or timedelta(days=settings.REFRESH_TOKEN_EXPIRE_DAYS))
    to_encode.update({"exp": expire, "refresh": True})
    return jwt.encode(to_encode, settings.JWT_SECRET_KEY, algorithm=settings.JWT_ALGORITHM)

def decode_token(token: str) -> Optional[dict]:
    try:
        payload = jwt.decode(token, settings.JWT_SECRET_KEY, algorithms=[settings.JWT_ALGORITHM])
        return payload
    except JWTError:
        return None""",

    "app/db/session.py": """from typing import Generator
from sqlmodel import SQLModel, create_engine, Session
from app.core.config import settings

engine = create_engine(settings.DATABASE_URL, echo=False)

def init_db() -> None:
    SQLModel.metadata.create_all(engine)

def get_session() -> Generator[Session, None, None]:
    with Session(engine) as session:
        yield session""",

    "app/models/all_models.py": """from uuid import UUID, uuid4
from datetime import datetime, timezone
from typing import Optional, List
from sqlmodel import SQLModel, Field, Relationship

class User(SQLModel, table=True):
    __tablename__ = "users"
    id: UUID = Field(default_factory=uuid4, primary_key=True)
    email: str = Field(unique=True, index=True, nullable=False)
    hashed_password: str = Field(nullable=False)
    created_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))

    documents: List["Document"] = Relationship(back_populates="user")

class Document(SQLModel, table=True):
    __tablename__ = "documents"
    id: UUID = Field(default_factory=uuid4, primary_key=True)
    user_id: UUID = Field(foreign_key="users.id", on_delete="CASCADE")
    filename: str = Field(nullable=False)
    storage_path: str = Field(nullable=False)
    extracted_text: str = Field(nullable=False)
    created_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))

    user: User = Relationship(back_populates="documents")
    resumos: List["Summary"] = Relationship(back_populates="document")
    flashcards: List["Flashcard"] = Relationship(back_populates="document")
    questions: List["Question"] = Relationship(back_populates="document")
    chat_history: List["ChatMessage"] = Relationship(back_populates="document")

class Summary(SQLModel, table=True):
    __tablename__ = "summaries"
    id: UUID = Field(default_factory=uuid4, primary_key=True)
    document_id: UUID = Field(foreign_key="documents.id", on_delete="CASCADE")
    level: str = Field(nullable=False)
    content: str = Field(nullable=False)
    created_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))

    document: Document = Relationship(back_populates="resumos")

class Flashcard(SQLModel, table=True):
    __tablename__ = "flashcards"
    id: UUID = Field(default_factory=uuid4, primary_key=True)
    document_id: UUID = Field(foreign_key="documents.id", on_delete="CASCADE")
    front: str = Field(nullable=False)
    back: str = Field(nullable=False)
    created_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))

    document: Document = Relationship(back_populates="flashcards")

class Question(SQLModel, table=True):
    __tablename__ = "questions"
    id: UUID = Field(default_factory=uuid4, primary_key=True)
    document_id: UUID = Field(foreign_key="documents.id", on_delete="CASCADE")
    statement: str = Field(nullable=False)
    options: str = Field(nullable=False)
    correct_option: str = Field(nullable=False)
    created_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))

    document: Document = Relationship(back_populates="questions")

class ChatMessage(SQLModel, table=True):
    __tablename__ = "chat_messages"
    id: UUID = Field(default_factory=uuid4, primary_key=True)
    document_id: UUID = Field(foreign_key="documents.id", on_delete="CASCADE")
    role: str = Field(nullable=False)
    content: str = Field(nullable=False)
    created_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))

    document: Document = Relationship(back_populates="chat_history")""",

    "app/schemas/all_schemas.py": """from pydantic import BaseModel, EmailStr, Field
from uuid import UUID
from datetime import datetime
from typing import List, Optional

class UserRegister(BaseModel):
    email: EmailStr
    password: str = Field(min_length=6)

class UserLogin(BaseModel):
    email: EmailStr
    password: str

class TokenResponse(BaseModel):
    access_token: str
    refresh_token: str
    token_type: str = "bearer"

class TokenRefreshRequest(BaseModel):
    refresh_token: str

class DocumentResponse(BaseModel):
    id: UUID
    filename: str
    created_at: datetime

class SummaryRequest(BaseModel):
    level: str = Field(pattern="^(1_min|5_min|complete)$")

class SummaryResponse(BaseModel):
    document_id: UUID
    level: str
    content: str

class FlashcardBase(BaseModel):
    front: str
    back: str

class FlashcardListResponse(BaseModel):
    flashcards: List[FlashcardBase]

class QuestionBase(BaseModel):
    statement: str
    options: List[str]
    correct_option: str

class QuestionListResponse(BaseModel):
    questions: List[QuestionBase]

class ChatRequest(BaseModel):
    message: str

class ChatResponse(BaseModel):
    answer: str

class TranslateRequest(BaseModel):
    language: str
    target_type: str = Field(pattern="^(text|summary)$")
    summary_level: Optional[str] = "1_min"

class TranslateResponse(BaseModel):
    translated_text: str""",

    "app/repositories/base_repository.py": """from typing import Generic, TypeVar, List, Optional
from uuid import UUID
from sqlmodel import Session, select

T = TypeVar("T")

class BaseRepository(Generic[T]):
    def __init__(self, session: Session, model: type[T]):
        self.session = session
        self.model = model

    def add(self, entity: T) -> T:
        self.session.add(entity)
        self.session.commit()
        self.session.refresh(entity)
        return entity

    def get_by_id(self, id: UUID) -> Optional[T]:
        return self.session.get(self.model, id)

    def list_all(self) -> List[T]:
        statement = select(self.model)
        return list(self.session.exec(statement).all())""",

    "app/repositories/specific_repositories.py": """from uuid import UUID
from typing import Optional, List
from sqlmodel import Session, select
from app.repositories.base_repository import BaseRepository
from app.models.all_models import User, Document, Summary, Flashcard, Question, ChatMessage

class UserRepository(BaseRepository[User]):
    def __init__(self, session: Session):
        super().__init__(session, User)

    def get_by_email(self, email: str) -> Optional[User]:
        return self.session.exec(select(User).where(User.email == email)).first()

class DocumentRepository(BaseRepository[Document]):
    def __init__(self, session: Session):
        super().__init__(session, Document)

    def get_user_document(self, document_id: UUID, user_id: UUID) -> Optional[Document]:
        return self.session.exec(
            select(Document).where(Document.id == document_id, Document.user_id == user_id)
        ).first()

    def list_by_user(self, user_id: UUID) -> List[Document]:
        return list(self.session.exec(select(Document).where(Document.user_id == user_id)).all())

class SummaryRepository(BaseRepository[Summary]):
    def __init__(self, session: Session):
        super().__init__(session, Summary)

    def get_by_doc_and_level(self, document_id: UUID, level: str) -> Optional[Summary]:
        return self.session.exec(
            select(Summary).where(Summary.document_id == document_id, Summary.level == level)
        ).first()

class FlashcardRepository(BaseRepository[Flashcard]):
    def __init__(self, session: Session):
        super().__init__(session, Flashcard)

    def get_by_document(self, document_id: UUID) -> List[Flashcard]:
        return list(self.session.exec(select(Flashcard).where(Flashcard.document_id == document_id)).all())

class QuestionRepository(BaseRepository[Question]):
    def __init__(self, session: Session):
        super().__init__(session, Question)

    def get_by_document(self, document_id: UUID) -> List[Question]:
        return list(self.session.exec(select(Question).where(Question.document_id == document_id)).all())

class ChatMessageRepository(BaseRepository[ChatMessage]):
    def __init__(self, session: Session):
        super().__init__(session, ChatMessage)

    def get_history(self, document_id: UUID, limit: int = 20) -> List[ChatMessage]:
        return list(self.session.exec(
            select(ChatMessage)
            .where(ChatMessage.document_id == document_id)
            .order_by(ChatMessage.created_at.asc())
            .limit(limit)
        ).all())""",

    "app/services/extractor_strategies.py": """import abc
import fitz
import docx
import easyocr
import numpy as np
from PIL import Image

class TextExtractionStrategy(abc.ABC):
    @abc.abstractmethod
    def extract(self, file_path: str) -> str:
        pass

class PdfExtractionStrategy(TextExtractionStrategy):
    def extract(self, file_path: str) -> str:
        text_content = []
        with fitz.open(file_path) as doc:
            for page in doc:
                text_content.append(page.get_text())
        extracted = "".join(text_content).strip()
        if not extracted:
            ocr_strategy = OcrExtractionStrategy()
            return ocr_strategy.extract(file_path)
        return extracted

class DocxExtractionStrategy(TextExtractionStrategy):
    def extract(self, file_path: str) -> str:
        doc = docx.Document(file_path)
        return "\\n".join([paragraph.text for paragraph in doc.paragraphs]).strip()

class OcrExtractionStrategy(TextExtractionStrategy):
    def __init__(self):
        self.reader = easyocr.Reader(["pt", "en"], gpu=False)

    def extract(self, file_path: str) -> str:
        if file_path.lower().endswith(".pdf"):
            text_content = []
            with fitz.open(file_path) as doc:
                for page in doc:
                    pix = page.get_pixmap()
                    img_bytes = pix.tobytes("png")
                    result = self.reader.readtext(img_bytes, detail=0)
                    text_content.append("\\n".join(result))
            return "\\n".join(text_content).strip()
        else:
            with Image.open(file_path) as img:
                img_np = np.array(img)
                result = self.reader.readtext(img_np, detail=0)
                return "\\n".join(result).strip()

class TextExtractorFactory:
    @staticmethod
    def get_strategy(file_extension: str) -> TextExtractionStrategy:
        ext = file_extension.lower()
        if ext == ".pdf":
            return PdfExtractionStrategy()
        elif ext in [".docx", ".doc"]:
            return DocxExtractionStrategy()
        elif ext in [".png", ".jpg", ".jpeg"]:
            return OcrExtractionStrategy()
        raise ValueError("Formato de arquivo não suportado.")""",

    "app/services/llm_service.py": """import json
from google import genai
from google.genai import types
from app.core.config import settings

class GeminiLlmService:
    def __init__(self):
        self.client = genai.Client(api_key=settings.GEMINI_API_KEY)
        self.model = "gemini-2.5-flash"

    def generate_summary(self, text: str, level: str) -> str:
        prompts = {
            "1_min": "Forneça um resumo ultra-condensado de no máximo 3 frases legível em 1 minuto.",
            "5_min": "Forneça um resumo detalhado e estruturado em seções principais legível em 5 minutos.",
            "complete": "Forneça um resumo exaustivo do seguinte documento."
        }
        prompt = f"{prompts.get(level)}\\n\\nDocumento:\\n{text}"
        response = self.client.models.generate_content(model=self.model, contents=prompt)
        return response.text

    def generate_flashcards(self, text: str) -> list:
        prompt = f"Gere flashcards em formato JSON válido (array de objetos com chaves 'front' e 'back') baseados no texto:\\n{text}"
        response = self.client.models.generate_content(
            model=self.model, contents=prompt, config=types.GenerateContentConfig(response_mime_type="application/json")
        )
        return json.loads(response.text)

    def generate_questions(self, text: str) -> list:
        prompt = f"Gere 5 questões de múltipla escolha em JSON válido (chaves 'statement', 'options', 'correct_option') baseadas no texto:\\n{text}"
        response = self.client.models.generate_content(
            model=self.model, contents=prompt, config=types.GenerateContentConfig(response_mime_type="application/json")
        )
        return json.loads(response.text)

    def translate_content(self, text: str, target_language: str) -> str:
        prompt = f"Traduza o seguinte texto para o idioma '{target_language}':\\n\\n{text}"
        response = self.client.models.generate_content(model=self.model, contents=prompt)
        return response.text

    def execute_chat_turn(self, history_tuples: list, current_message: str, document_context: str) -> str:
        chat_contents = [f"Contexto do documento:\\n{document_context}\\n\\nInstrução: Responda baseando-se estritamente no contexto."]
        for role, content in history_tuples:
            chat_contents.append(f"{role}: {content}")
        chat_contents.append(f"user: {current_message}")
        response = self.client.models.generate_content(model=self.model, contents="\\n".join(chat_contents))
        return response.text""",

    "app/services/document_service.py": """import os
import uuid
from uuid import UUID
from fastapi import UploadFile, HTTPException, status
from fastapi.concurrency import run_in_threadpool
from app.core.config import settings
from app.models.all_models import Document, Summary, Flashcard, Question, ChatMessage
from app.repositories.specific_repositories import DocumentRepository, SummaryRepository, FlashcardRepository, QuestionRepository, ChatMessageRepository
from app.services.extractor_strategies import TextExtractorFactory
from app.services.llm_service import GeminiLlmService
import json

class DocumentService:
    def __init__(self, doc_repo: DocumentRepository, sum_repo: SummaryRepository, flash_repo: FlashcardRepository, quest_repo: QuestionRepository, chat_repo: ChatMessageRepository, llm: GeminiLlmService):
        self.doc_repo = doc_repo
        self.sum_repo = sum_repo
        self.flash_repo = flash_repo
        self.quest_repo = quest_repo
        self.chat_repo = chat_repo
        self.llm = llm

    async def validate_and_save_file(self, file: UploadFile) -> tuple[str, str]:
        os.makedirs(settings.UPLOAD_DIR, exist_ok=True)
        header = await file.read(4)
        await file.seek(0)
        is_pdf = header == b'%PDF'
        is_docx = header == b'PK\\x03\\x04'
        is_png = header == b'\\x89PNG'
        is_jpg = header[:3] == b'\\xff\\xd8\\xff'

        if not (is_pdf or is_docx or is_png or is_jpg):
            raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Arquivo inválido.")

        file_ext = os.path.splitext(file.filename)[1]
        unique_filename = f"{uuid.uuid4()}{file_ext}"
        storage_path = os.path.join(settings.UPLOAD_DIR, unique_filename)
        with open(storage_path, "wb") as buffer:
            while chunk := await file.read(1024 * 64):
                buffer.write(chunk)
        return unique_filename, storage_path

    async def process_document_upload(self, file: UploadFile, user_id: UUID) -> Document:
        filename, storage_path = await self.validate_and_save_file(file)
        file_ext = os.path.splitext(filename)[1]
        try:
            extractor = TextExtractorFactory.get_strategy(file_ext)
            extracted_text = await run_in_threadpool(extractor.extract, storage_path)
        except Exception:
            if os.path.exists(storage_path):
                os.remove(storage_path)
            raise HTTPException(status_code=status.HTTP_422_UNPROCESSABLE_ENTITY, detail="Erro de extração.")
        doc = Document(user_id=user_id, filename=file.filename, storage_path=storage_path, extracted_text=extracted_text)
        return self.doc_repo.add(doc)

    async def get_summary(self, document_id: UUID, user_id: UUID, level: str) -> Summary:
        doc = self.doc_repo.get_user_document(document_id, user_id)
        if not doc: raise HTTPException(status_code=404, detail="Não encontrado.")
        cached = self.sum_repo.get_by_doc_and_level(document_id, level)
        if cached: return cached
        summary_text = await run_in_threadpool(self.llm.generate_summary, doc.extracted_text, level)
        return self.sum_repo.add(Summary(document_id=document_id, level=level, content=summary_text))

    async def get_flashcards(self, document_id: UUID, user_id: UUID) -> list[Flashcard]:
        doc = self.doc_repo.get_user_document(document_id, user_id)
        if not doc: raise HTTPException(status_code=404, detail="Não encontrado.")
        existing = self.flash_repo.get_by_document(document_id)
        if existing: return existing
        cards_data = await run_in_threadpool(self.llm.generate_flashcards, doc.extracted_text)
        return [self.flash_repo.add(Flashcard(document_id=document_id, front=i["front"], back=i["back"])) for i in cards_data]

    async def get_questions(self, document_id: UUID, user_id: UUID) -> list[Question]:
        doc = self.doc_repo.get_user_document(document_id, user_id)
        if not doc: raise HTTPException(status_code=404, detail="Não encontrado.")
        existing = self.quest_repo.get_by_document(document_id)
        if existing: return existing
        q_data = await run_in_threadpool(self.llm.generate_questions, doc.extracted_text)
        return [self.quest_repo.add(Question(document_id=document_id, statement=i["statement"], options=json.dumps(i["options"]), correct_option=i["correct_option"])) for i in q_data]

    async def interact_in_chat(self, document_id: UUID, user_id: UUID, user_message: str) -> str:
        doc = self.doc_repo.get_user_document(document_id, user_id)
        if not doc: raise HTTPException(status_code=404, detail="Não encontrado.")
        history = self.chat_repo.get_history(document_id)
        h_tuples = [(msg.role, msg.content) for msg in history]
        answer = await run_in_threadpool(self.llm.execute_chat_turn, h_tuples, user_message, doc.extracted_text)
        self.chat_repo.add(ChatMessage(document_id=document_id, role="user", content=user_message))
        self.chat_repo.add(ChatMessage(document_id=document_id, role="model", content=answer))
        return answer

    async def translate_document_element(self, document_id: UUID, user_id: UUID, language: str, target_type: str, summary_level: str) -> str:
        doc = self.doc_repo.get_user_document(document_id, user_id)
        if not doc: raise HTTPException(status_code=404, detail="Não encontrado.")
        text = doc.extracted_text if target_type == "text" else (await self.get_summary(document_id, user_id, summary_level)).content
        return await run_in_threadpool(self.llm.translate_content, text, language)""",

    "app/services/auth_service.py": """from fastapi import HTTPException, status
from app.repositories.specific_repositories import UserRepository
from app.schemas.all_schemas import UserRegister, UserLogin, TokenResponse
from app.models.all_models import User
from app.core.security import get_password_hash, verify_password, create_access_token, create_refresh_token, decode_token

class AuthService:
    def __init__(self, user_repo: UserRepository):
        self.user_repo = user_repo

    def register_user(self, schema: UserRegister) -> User:
        if self.user_repo.get_by_email(schema.email):
            raise HTTPException(status_code=400, detail="Email já cadastrado.")
        user = User(email=schema.email, hashed_password=get_password_hash(schema.password))
        return self.user_repo.add(user)

    def authenticate_user(self, schema: UserLogin) -> TokenResponse:
        user = self.user_repo.get_by_email(schema.email)
        if not user or not verify_password(schema.password, user.hashed_password):
            raise HTTPException(status_code=401, detail="Credenciais incorretas.")
        return TokenResponse(access_token=create_access_token({"sub": str(user.id)}), refresh_token=create_refresh_token({"sub": str(user.id)}))

    def refresh_session(self, refresh_token: str) -> TokenResponse:
        payload = decode_token(refresh_token)
        if not payload or not payload.get("refresh"):
            raise HTTPException(status_code=401, detail="Token inválido.")
        uid = payload.get("sub")
        return TokenResponse(access_token=create_access_token({"sub": uid}), refresh_token=create_refresh_token({"sub": uid}))""",

    "app/api/dependencies.py": """from fastapi import Depends, HTTPException, status
from fastapi.security import OAuth2PasswordBearer
from sqlmodel import Session
from app.db.session import get_session
from app.core.security import decode_token
from app.repositories.specific_repositories import UserRepository, DocumentRepository, SummaryRepository, FlashcardRepository, QuestionRepository, ChatMessageRepository
from app.services.auth_service import AuthService
from app.services.document_service import DocumentService
from app.services.llm_service import GeminiLlmService
from uuid import UUID

oauth2_scheme = OAuth2PasswordBearer(tokenUrl="api/v1/auth/login")

def get_user_repo(session: Session = Depends(get_session)) -> UserRepository:
    return UserRepository(session)

def get_auth_service(user_repo: UserRepository = Depends(get_user_repo)) -> AuthService:
    return AuthService(user_repo)

def get_current_user_id(token: str = Depends(oauth2_scheme)) -> UUID:
    payload = decode_token(token)
    if not payload or payload.get("refresh"):
        raise HTTPException(status_code=401, detail="Token inválido.")
    return UUID(payload.get("sub"))

def get_document_service(session: Session = Depends(get_session)) -> DocumentService:
    return DocumentService(
        doc_repo=DocumentRepository(session), sum_repo=SummaryRepository(session),
        flash_repo=FlashcardRepository(session), quest_repo=QuestionRepository(session),
        chat_repo=ChatMessageRepository(session), llm=GeminiLlmService()
    )""",

    "app/api/v1/endpoints/auth.py": """from fastapi import APIRouter, Depends
from app.schemas.all_schemas import UserRegister, UserLogin, TokenResponse, TokenRefreshRequest
from app.services.auth_service import AuthService
from app.api.dependencies import get_auth_service

router = APIRouter(prefix="/auth", tags=["Autenticação"])

@router.post("/register", status_code=201)
def register(data: UserRegister, service: AuthService = Depends(get_auth_service)):
    service.register_user(data)
    return {"message": "Registrado com sucesso."}

@router.post("/login", response_model=TokenResponse)
def login(data: UserLogin, service: AuthService = Depends(get_auth_service)):
    return service.authenticate_user(data)

@router.post("/refresh", response_model=TokenResponse)
def refresh(data: TokenRefreshRequest, service: AuthService = Depends(get_auth_service)):
    return service.refresh_session(data.refresh_token)""",

    "app/api/v1/endpoints/documents.py": """from fastapi import APIRouter, Depends, UploadFile, File, HTTPException, status
from typing import List
from uuid import UUID
import json
from app.schemas.all_schemas import (
    DocumentResponse, SummaryRequest, SummaryResponse, 
    FlashcardListResponse, QuestionListResponse, ChatRequest, 
    ChatResponse, TranslateRequest, TranslateResponse
)
from app.services.document_service import DocumentService
from app.api.dependencies import get_document_service, get_current_user_id
from app.core.config import settings

router = APIRouter(prefix="/documents", tags=["Documentos & IA"])

@router.post("/upload", response_model=DocumentResponse, status_code=201)
async def upload(file: UploadFile = File(...), user_id: UUID = Depends(get_current_user_id), service: DocumentService = Depends(get_document_service)):
    file.file.seek(0, 2)
    fs = file.file.tell()
    await file.seek(0)
    if fs > settings.MAX_FILE_SIZE_BYTES: raise HTTPException(status_code=413, detail="Arquivo grande.")
    return await service.process_document_upload(file, user_id)

@router.post("/{document_id}/summary", response_model=SummaryResponse)
async def summarize(document_id: UUID, data: SummaryRequest, user_id: UUID = Depends(get_current_user_id), service: DocumentService = Depends(get_document_service)):
    s = await service.get_summary(document_id, user_id, data.level)
    return SummaryResponse(document_id=s.document_id, level=s.level, content=s.content)

@router.get("/{document_id}/flashcards", response_model=FlashcardListResponse)
async def flashcards(document_id: UUID, user_id: UUID = Depends(get_current_user_id), service: DocumentService = Depends(get_document_service)):
    c = await service.get_flashcards(document_id, user_id)
    return {"flashcards": [{"front": i.front, "back": i.back} for i in c]}

@router.get("/{document_id}/questions", response_model=QuestionListResponse)
async def questions(document_id: UUID, user_id: UUID = Depends(get_current_user_id), service: DocumentService = Depends(get_document_service)):
    q = await service.get_questions(document_id, user_id)
    return {"questions": [{"statement": i.statement, "options": json.loads(i.options), "correct_option": i.correct_option} for i in q]}

@router.post("/{document_id}/chat", response_model=ChatResponse)
async def chat(document_id: UUID, data: ChatRequest, user_id: UUID = Depends(get_current_user_id), service: DocumentService = Depends(get_document_service)):
    a = await service.interact_in_chat(document_id, user_id, data.message)
    return ChatResponse(answer=a)

@router.post("/{document_id}/translate", response_model=TranslateResponse)
async def translate(document_id: UUID, data: TranslateRequest, user_id: UUID = Depends(get_current_user_id), service: DocumentService = Depends(get_document_service)):
    r = await service.translate_document_element(document_id, user_id, data.language, data.target_type, data.summary_level)
    return TranslateResponse(translated_text=r)

@router.get("/{document_id}/export")
async def export_data(document_id: UUID, user_id: UUID = Depends(get_current_user_id), service: DocumentService = Depends(get_document_service)):
    c = await service.get_flashcards(document_id, user_id)
    s = await service.get_summary(document_id, user_id, "1_min")
    return {"document_id": str(document_id), "cached_summary_1m": s.content, "flashcards": [{"front": i.front, "back": i.back} for i in c]}""",

    "app/main.py": """from fastapi import FastAPI, Request, status
from fastapi.responses import JSONResponse
from app.api.v1.endpoints import auth, documents
from app.db.session import init_db
import logging

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger("LeitorIA")

app = FastAPI(title="Leitor IA API Backend", version="1.0.0")

@app.on_event("startup")
def on_startup():
    init_db()

@app.exception_handler(Exception)
async def generic_exception_handler(request: Request, exc: Exception):
    logger.error(f"Erro Crítico: {str(exc)}", exc_info=True)
    return JSONResponse(status_code=500, content={"detail": "Erro interno no servidor."})

app.include_router(auth.router, prefix="/api/v1")
app.include_router(documents.router, prefix="/api/v1")"""
}

# Cria as pastas e os arquivos contendo o código de forma automatizada
print("🚀 Iniciando a extração do esqueleto da Clean Architecture...")
for filepath, content in project_files.items():
    dir_name = os.path.dirname(filepath)
    if dir_name:
        os.makedirs(dir_name, exist_ok=True)
    
    with open(filepath, "w", encoding="utf-8") as f:
        f.write(content.strip())
    print(f"  [Criado] -> {filepath}")

print("\n✨ Estrutura gerada com sucesso! Todos os arquivos estão nos seus devidos lugares.")